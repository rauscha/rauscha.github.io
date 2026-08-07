#!/usr/bin/env python3
"""Build a styled .docx and .pdf academic CV from a Markdown source.
Vendored from https://github.com/rauscha/md-cv — sync manually on material changes.

Canonical home: https://github.com/rauscha/md-cv
Usage: python build_cv.py path/to/CV.md [-o OUTDIR]

Format: `# Name` then contact lines; `## SECTION`; `### (a) Subsection`;
dated entries as `2018-2021 | description` (must start with a 4-digit year);
everything else is a plain paragraph. `**bold**` / `*italic*` inline.
Lines containing [TBC] stay in the source but are stripped from output.
"""
from __future__ import annotations

import argparse
import difflib
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Emu, Pt

# ---------------------------------------------------------------------------
# STYLE CONSTANTS — the entire visual design lives here. "Light refresh" =
# conservative academic layout, Calibri, tidy consistent spacing.
# ---------------------------------------------------------------------------
FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
NAME_SIZE = Pt(16)
INDENT = Emu(914400)            # 1" hanging indent + tab stop for dated entries
SPACE_BEFORE_PARA = Pt(0)
SPACE_BEFORE_SECTION = Pt(12)
SPACE_AFTER_SECTION = Pt(4)
SPACE_BEFORE_SUBSECTION = Pt(6)
SPACE_AFTER_PARA = Pt(2)
SPACE_GAP = Pt(8)           # extra space_after on a paragraph followed by a blank line in the source


@dataclass
class Entry:
    dates: str
    text: str
    gap: bool = False


@dataclass
class Para:
    text: str
    gap: bool = False


@dataclass
class Subsection:
    title: str
    items: list = field(default_factory=list)


@dataclass
class Section:
    title: str
    items: list = field(default_factory=list)


@dataclass
class CV:
    name: str
    contact: list
    sections: list


DATE_ENTRY_RE = re.compile(r"^((?:\d{4}|Current)[^|]*)\|(.*)$")


def parse_cv(text: str) -> CV:
    name = ""
    contact: list[str] = []
    sections: list[Section] = []
    section: Section | None = None
    sub: Subsection | None = None
    saw_blank = False
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            saw_blank = True
            continue
        if "[TBC]" in line:
            continue
        if line.startswith("### "):
            if section is None:
                raise ValueError(
                    f"subsection before any '## SECTION' header: {line!r}"
                )
            sub = Subsection(line[4:].strip())
            section.items.append(sub)
            saw_blank = False
        elif line.startswith("## "):
            section = Section(line[3:].strip())
            sub = None
            sections.append(section)
            saw_blank = False
        elif line.startswith("# "):
            name = line[2:].strip()
            saw_blank = False
        else:
            m = DATE_ENTRY_RE.match(line)
            item = Entry(m.group(1).strip(), m.group(2).strip()) if m else Para(line)
            if sub is not None:
                target = sub.items
            elif section is not None:
                target = section.items
            else:
                target = None
            if target is not None:
                if saw_blank and target:
                    target[-1].gap = True
                target.append(item)
            else:
                contact.append(line)
            saw_blank = False
    return CV(name, contact, sections)


TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|(?<![\w*])\*[^*\s][^*]*\*(?![\w*]))")


def tokenize_runs(text: str) -> list:
    out = []
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True, False))
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            out.append((part[1:-1], False, True))
        else:
            out.append((part, False, False))
    return out


def _add_runs(p, text: str):
    for t, bold, italic in tokenize_runs(text):
        r = p.add_run(t)
        r.bold = bold
        r.italic = italic
    return p


def _render_items(doc, items):
    for item in items:
        if isinstance(item, Subsection):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = SPACE_BEFORE_SUBSECTION
            p.add_run(item.title).bold = True
            if item.items:
                _render_items(doc, item.items)
            else:
                doc.add_paragraph("None")
        elif isinstance(item, Entry):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = INDENT
            pf.first_line_indent = -INDENT
            pf.tab_stops.add_tab_stop(INDENT, WD_TAB_ALIGNMENT.LEFT)
            p.add_run(item.dates + "\t")
            _add_runs(p, item.text)
            if item.gap:
                pf.space_after = SPACE_GAP
        else:
            p = doc.add_paragraph()
            _add_runs(p, item.text)
            if item.gap:
                p.paragraph_format.space_after = SPACE_GAP


def render_docx(cv: CV, path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_before = SPACE_BEFORE_PARA
    normal.paragraph_format.space_after = SPACE_AFTER_PARA

    p = doc.add_paragraph()
    r = p.add_run(cv.name)
    r.bold = True
    r.font.size = NAME_SIZE

    for line in cv.contact:
        _add_runs(doc.add_paragraph(), line)

    for section in cv.sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = SPACE_BEFORE_SECTION
        p.paragraph_format.space_after = SPACE_AFTER_SECTION
        p.add_run(section.title).bold = True
        _render_items(doc, section.items)

    doc.save(str(path))


def _export_via_word(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client
    except ImportError:
        return False
    try:
        word = win32com.client.DispatchEx("Word.Application")
    except Exception:
        return False
    try:
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        try:
            doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)  # wdFormatPDF
        finally:
            doc.Close(False)
            del doc  # release proxy while Word is still alive (avoids RPC warnings)
    except Exception as exc:
        print(f"Word export failed ({exc}); trying LibreOffice...", file=sys.stderr)
        return False
    finally:
        word.Quit()
        del word
    return True


def _export_via_soffice(docx_path: Path, pdf_path: Path) -> bool:
    soffice = shutil.which("soffice")
    if not soffice:
        for candidate in (
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path("/usr/bin/soffice"),
        ):
            if candidate.exists():
                soffice = str(candidate)
                break
    if not soffice:
        return False
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", str(pdf_path.parent), str(docx_path)],
            check=True, timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        print(f"LibreOffice export failed ({exc})", file=sys.stderr)
        return False
    produced = pdf_path.parent / (docx_path.stem + ".pdf")
    if produced != pdf_path:
        produced.replace(pdf_path)
    return pdf_path.exists()


def export_pdf(docx_path: Path, pdf_path: Path) -> str:
    if _export_via_word(docx_path, pdf_path):
        return "word"
    if _export_via_soffice(docx_path, pdf_path):
        return "libreoffice"
    raise RuntimeError(
        "No PDF exporter found. Install Microsoft Word (plus `pip install pywin32`) "
        "or LibreOffice."
    )


def extract_pdf_text(pdf_path: Path):
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return text, len(reader.pages)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Build a .docx and .pdf CV from Markdown.")
    ap.add_argument("markdown", type=Path)
    ap.add_argument("-o", "--outdir", type=Path, default=None)
    args = ap.parse_args(argv)

    outdir = args.outdir or args.markdown.parent
    outdir.mkdir(parents=True, exist_ok=True)
    docx_path = outdir / (args.markdown.stem + ".docx")
    pdf_path = outdir / (args.markdown.stem + ".pdf")

    old_text = extract_pdf_text(pdf_path)[0] if pdf_path.exists() else None

    cv = parse_cv(args.markdown.read_text(encoding="utf-8"))
    render_docx(cv, docx_path)
    engine = export_pdf(docx_path, pdf_path)

    new_text, pages = extract_pdf_text(pdf_path)
    if "TBC" in new_text:
        sys.exit("ERROR: [TBC] content leaked into the PDF")
    print(f"Built {docx_path.name} and {pdf_path.name} via {engine}; {pages} page(s).")
    if old_text is not None:
        diff = list(difflib.unified_diff(
            old_text.splitlines(), new_text.splitlines(),
            "previous.pdf", "new.pdf", lineterm=""))
        print("\n".join(diff) if diff else "No text changes vs previous PDF.")


if __name__ == "__main__":
    main()
